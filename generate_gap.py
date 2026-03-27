from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
for section in doc.sections:
    section.top_margin = Cm(2); section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5); section.right_margin = Cm(2.5)

# ── helpers ──────────────────────────────────────────────────────
def heading1(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(16); r.font.name = "Calibri"
    r.font.color.rgb = RGBColor(0, 70, 127)
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1');    bot.set(qn('w:color'), '00467F')
    pBdr.append(bot); pPr.append(pBdr)

def heading2(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.size = Pt(13); r.font.name = "Calibri"
    r.font.color.rgb = RGBColor(31, 73, 125)
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)

def body(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(11); r.font.name = "Calibri"
    if color: r.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(4)

def bullet(text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.size = Pt(11); r.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(2)

def colorbox(text, fill_hex, text_color=(0,0,0)):
    tbl = doc.add_table(rows=1, cols=1); tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]; cell.text = text
    r = cell.paragraphs[0].runs[0]
    r.font.size = Pt(11); r.font.name = "Calibri"
    r.font.color.rgb = RGBColor(*text_color)
    tc = cell._tc; tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd); doc.add_paragraph()

def add_table(headers, rows, hdr_color='1F497D', alt='DCE6F1'):
    tbl = doc.add_table(rows=1+len(rows), cols=len(headers))
    tbl.style = 'Table Grid'; tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = tbl.rows[0].cells[i]; c.text = h
        r = c.paragraphs[0].runs[0]
        r.bold=True; r.font.size=Pt(10); r.font.name="Calibri"
        r.font.color.rgb = RGBColor(255,255,255)
        c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc=c._tc; tcPr=tc.get_or_add_tcPr()
        shd=OxmlElement('w:shd'); shd.set(qn('w:val'),'clear')
        shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hdr_color); tcPr.append(shd)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = tbl.rows[ri+1].cells[ci]; c.text = val
            r = c.paragraphs[0].runs[0]; r.font.size=Pt(10); r.font.name="Calibri"
            if ri%2==0:
                tc=c._tc; tcPr=tc.get_or_add_tcPr()
                shd=OxmlElement('w:shd'); shd.set(qn('w:val'),'clear')
                shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),alt); tcPr.append(shd)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════
p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run("Gap Analysis — Finomnia AI in PEF V05 vs Metis ACE")
r.bold=True; r.font.size=Pt(22); r.font.name="Calibri"; r.font.color.rgb=RGBColor(0,70,127)

p2 = doc.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Funzionalità del documento Finomnia: cosa è già implementato in Metis")
r2.italic=True; r2.font.size=Pt(13); r2.font.name="Calibri"; r2.font.color.rgb=RGBColor(89,89,89)

doc.add_paragraph()
p3 = doc.add_paragraph(); p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Documento interno  |  Marzo 2026  |  Confidenziale")
r3.font.size=Pt(11); r3.font.name="Calibri"; r3.font.color.rgb=RGBColor(128,128,128)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 1. SINTESI ESECUTIVA
# ══════════════════════════════════════════════════════════════════
heading1("1. Sintesi Esecutiva")

colorbox(
    "✅ RISULTATO CHIAVE: Tutti gli 8 casi d'uso descritti nel documento Finomnia "
    "AI in PEF V05 (19/03/2026) sono stati implementati in Metis ACE come M1–M8.\n\n"
    "Il documento Finomnia è di fatto la specifica funzionale originale da cui Metis è stato sviluppato. "
    "Metis copre il 100% dei requisiti core e aggiunge funzionalità non previste nel documento originale.",
    fill_hex='D9EAD3', text_color=(0,80,0)
)

add_table(
    ["Dimensione", "Valore"],
    [
        ["Casi d'uso Finomnia (core)",     "8"],
        ["Implementati in Metis",           "8  (100%)"],
        ["Gap sui requisiti core",          "0"],
        ["Evoluzioni future (non core)",    "7  (non ancora implementate — roadmap)"],
        ["Funzionalità extra di Metis",     "3  (Visual Policy Builder, Fraud Detection, Audit Trail XAI)"],
    ]
)

# ══════════════════════════════════════════════════════════════════
# 2. IL DOCUMENTO FINOMNIA — CONTESTO
# ══════════════════════════════════════════════════════════════════
heading1("2. Il Documento Finomnia AI in PEF V05 — Contesto")
body(
    "Il documento Finomnia 'AI in PEF V05' (versione del 19/03/2026) è una specifica funzionale "
    "che propone l'integrazione di 8 moduli AI nel processo PEF delle banche italiane. "
    "Si basa su tre fonti accademiche/normative:"
)
bullet("Paper Banca d'Italia — 'Credit Risk Assessment with Stacked Machine Learning' (Gennaio 2026)")
bullet("Accenture — 'The Age of AI: Banking's New Reality' (73% del tempo lavorativo bancario influenzabile dall'AI)")
bullet("EU Artificial Intelligence Act — Regolamento UE 2024/1689 (sistemi di credit scoring = Alto Rischio)")

body(
    "La filosofia del documento: trasformare l'analisi del credito da una 'fotografia del passato' "
    "(bilanci vecchi) a un sistema di monitoraggio dinamico e trasparente, in linea con "
    "il Codice della Crisi d'Impresa e i dettami EBA.",
    italic=True
)

# ══════════════════════════════════════════════════════════════════
# 3. GAP ANALYSIS — MODULI CORE (M1–M8)
# ══════════════════════════════════════════════════════════════════
heading1("3. Gap Analysis — 8 Moduli Core")

body("Confronto puntuale tra quanto descritto nel documento Finomnia e quanto implementato in Metis ACE:")

doc.add_paragraph()

def module_row(num, finomnia_desc, metis_impl, status, note=""):
    heading2(f"M{num} — {finomnia_desc[:60]}")
    add_table(
        ["Campo", "Finomnia V05 (Specifica)", "Metis ACE (Implementazione)", "Stato"],
        [[
            f"M{num}",
            finomnia_desc,
            metis_impl,
            status
        ]],
        hdr_color='1F497D'
    )
    if note:
        body(f"📌 Nota: {note}", italic=True, color=(70,70,70))

# M1
module_row(1,
    "L'IA rielabora i commenti delle ultime due revisioni PEF generando tre paragrafi distinti: "
    "evoluzione societaria, dinamica reddituale-patrimoniale storica, comportamento pregresso in CR.",
    "Implementato. LLM (Claude Sonnet) genera sintesi_societaria, sintesi_bilancio, sintesi_cr. "
    "Anti-hallucination check sui numeri. Source attribution XAI su ogni paragrafo.",
    "✅ IMPLEMENTATO",
    "Metis aggiunge il controllo anti-allucinazione e il confidence score per paragrafo — "
    "funzionalità non previste nel documento Finomnia."
)

# M2
module_row(2,
    "Il sistema scansiona fonti autorevoli (Sole 24Ore, ANSA) per descrivere attività e reputazione "
    "del soggetto e dei correlati, calcolando sentiment Positivo/Neutro/Negativo.",
    "Implementato. Web scraper con whitelist (Sole24Ore, ANSA, Reuters, Registro Imprese). "
    "FinBERT per sentiment analysis finanziario. Aggregazione pesata per autorevolezza fonte.",
    "✅ IMPLEMENTATO",
    "Metis aggiunge il modello FinBERT specializzato per italiano finanziario e l'identificazione "
    "automatica dei soggetti correlati emersi dalle notizie."
)

# M3
module_row(3,
    "L'IA calcola i principali indicatori di redditività, struttura e sostenibilità (EBITDA Margin, "
    "Leverage) confrontando Tzero con T-1 per identificare il trend interno.",
    "Implementato. 9 KPI deterministici (EBITDA Margin, ROE, ROA, Leverage D/E, Current Ratio, "
    "Interest Coverage, NFP/EBITDA, DSO, DPO). XGBoost risk scorer + SHAP.",
    "✅ IMPLEMENTATO",
    "Metis aggiunge 5 KPI non citati esplicitamente nel documento (ROE, ROA, Current Ratio, DSO, DPO) "
    "e il risk score XGBoost con spiegazione SHAP."
)

# M4
module_row(4,
    "Il sistema confronta i KPI aziendali con le medie di settore ISTAT per definire il "
    "posizionamento competitivo (Outperformer/Underperformer) e genera una matrice di valutazione sintetica.",
    "Implementato. Database ISTAT con mediane e percentili 25°/75° per ~200 codici ATECO. "
    "Matching automatico. Posizionamento percentile per ogni KPI.",
    "✅ IMPLEMENTATO",
    "Metis aggiunge il calcolo del percentile esatto (non solo above/below median) e "
    "l'output viene alimentato automaticamente in M7 e M8."
)

# M5
module_row(5,
    "L'IA confronta l'ultimo dato CR con la media storica (12 mesi) per rilevare tensioni "
    "finanziarie, distinguendo tra anomalie transitorie sanate e insoluti strutturali.",
    "Implementato. Time-series analysis su PostgreSQL/TimescaleDB. Z-score anomaly detection (soglia >2σ). "
    "Classificazione automatica TRANSITORIA vs STRUTTURALE (rientro entro 60 giorni).",
    "✅ IMPLEMENTATO",
    "Metis aggiunge il calcolo del tasso di utilizzo del credito (accordato vs utilizzato) "
    "con alert sopra 85%, non esplicitato nel documento Finomnia."
)

# M6
module_row(6,
    "Verifica di integrità (Cross-Check) che calcola il Delta percentuale tra i debiti bancari "
    "a bilancio e l'utilizzato in CR, generando alert se lo scostamento supera il 10%.",
    "Implementato. Calcolo delta% con soglie OK (<10%), WARNING (10–25%), CRITICAL (>25%). "
    "Analisi automatica delle possibili cause. Raccomandazione all'analista.",
    "✅ IMPLEMENTATO",
    "Metis aggiunge l'analisi automatica delle cause (sfasamento temporale, debiti non censiti in CR, "
    "utilizzi post-chiusura bilancio) — non prevista nel documento Finomnia."
)

# M7
module_row(7,
    "Stima della sostenibilità futura del debito tramite un DSCR prospettico a 12 mesi, "
    "basato su scenari (Ottimistico/Base/Stress) selezionati dall'IA in funzione ai trend CR e bilancio.",
    "Implementato. Formula DSCR = (EBITDA proiettato - Tasse - Capex) / (Quota capitale + Interessi). "
    "3 scenari con parametri configurabili. Selezione automatica scenario da output M3 e M5.",
    "✅ IMPLEMENTATO",
    "Metis aggiunge la selezione automatica dello scenario basata sugli output di M3 e M5 "
    "e parametri configurabili per ogni scenario (revenue growth, rate delta, cost growth)."
)

# M8
module_row(8,
    "Generazione di una matrice strategica 2x2 che mappa automaticamente i dati di bilancio "
    "(Forza/Debolezza) e il contesto geografico-settoriale (Opportunità/Minacce) su quadranti predefiniti.",
    "Implementato. Rule engine + LLM narrative. Ogni item SWOT linkato al modulo sorgente "
    "(M3, M4, M5, M2). Narrativa LLM di sintesi executive dei 4 quadranti.",
    "✅ IMPLEMENTATO",
    "Metis aggiunge il tracciamento del modulo sorgente (source_module) per ogni item SWOT "
    "e la narrativa LLM executive — non previsti nel documento Finomnia."
)

# ══════════════════════════════════════════════════════════════════
# 4. EVOLUZIONI FUTURE — NON ANCORA IMPLEMENTATE
# ══════════════════════════════════════════════════════════════════
heading1("4. Evoluzioni Future — Non Ancora Implementate (Roadmap)")

body(
    "Il documento Finomnia (pagine 11–12) elenca 7 evoluzioni future proposte oltre ai moduli core. "
    "Nessuna di queste è attualmente implementata in Metis. Rappresentano la roadmap di sviluppo:"
)

add_table(
    ["#", "Evoluzione Proposta (Finomnia V05)", "Priorità Stimata", "Complessità", "Stato Metis"],
    [
        ["E1", "Simulatore prodotto finanziario: individua il prodotto più profittevole per il cliente "
               "(es. linea p-endo vs p-uto, con/senza garanzie, con/senza notifica)",
               "Alta", "Alta", "❌ Non implementato"],
        ["E2", "Indicatori di pagamento per cluster: verifica se il debitore performa in linea "
               "con il proprio cluster di appartenenza",
               "Media", "Media", "❌ Non implementato"],
        ["E3", "Teoria dei giochi in PEF: strumenti decisionali sofisticati per scenari "
               "multi-agente nella negoziazione del credito",
               "Bassa", "Molto Alta", "❌ Non implementato"],
        ["E4", "Monitoraggio basato su analisi predittiva: early warning su portafoglio",
               "Alta", "Alta", "❌ Non implementato"],
        ["E5", "Indicatori per la valutazione del rischio AR (Atipicità/Rilevanza)",
               "Media", "Media", "❌ Non implementato"],
        ["E6", "Individuazione cluster omogenei di clientela per early warning comportamentale",
               "Alta", "Alta", "❌ Non implementato"],
        ["E7", "Lettura e sintesi allegati documentali: OCR + AI per accelerare l'analisi "
               "dei documenti allegati alla pratica",
               "Alta", "Media", "❌ Non implementato"],
        ["E8", "Analisi aderenza pratica alle policy creditizie della banca: verifica automatica "
               "della conformità della pratica alle regole interne",
               "Media", "Media", "❌ Non implementato"],
    ]
)

# ══════════════════════════════════════════════════════════════════
# 5. FUNZIONALITÀ EXTRA DI METIS (NON NEL DOCUMENTO FINOMNIA)
# ══════════════════════════════════════════════════════════════════
heading1("5. Funzionalità Extra di Metis — Non Previste nel Documento Finomnia")

body("Metis ha implementato 3 aree funzionali importanti non descritte nel documento Finomnia:")

add_table(
    ["Funzionalità Extra Metis", "Descrizione", "Valore Aggiunto"],
    [
        ["Visual Policy Builder",
         "Editor drag & drop no-code per configurare il flusso decisionale del credito. "
         "Blocchi: Data Ingestion, Fraud Detection, AI Scoring, Auto-Approve, Reject.",
         "Personalizzazione senza codice. La banca adatta le soglie e le regole. Versionato."],
        ["Audit Trail Immutabile (EU AI Act)",
         "Ogni azione nel sistema genera un record immutabile con hash SHA-256 di input e output. "
         "Override analista obbligatoriamente motivato. Esportabile.",
         "Compliance EU AI Act nativa. Difendibile in sede di ispezione regolamentare."],
        ["Fraud Detection (Zest Protect-style)",
         "ML Anomaly Detection su documenti OCR. False Positive Limit configurabile (es. 4.3%). "
         "Integrato come blocco nel Visual Policy Builder.",
         "Riduce il rischio di frodi documentali prima ancora dell'analisi creditizia."],
    ]
)

# ══════════════════════════════════════════════════════════════════
# 6. RACCOMANDAZIONI PRIORITÀ ROADMAP
# ══════════════════════════════════════════════════════════════════
heading1("6. Raccomandazioni — Priorità di Sviluppo Roadmap")

body("Sulla base dell'analisi del documento Finomnia e del mercato, le evoluzioni da sviluppare per prime:")

add_table(
    ["Priorità", "Evoluzione", "Motivazione"],
    [
        ["🔴 1°", "E7 — Lettura allegati documentali (OCR + AI)",
         "Alto valore percepito, complessità tecnica media. Direttamente nella pipeline esistente."],
        ["🔴 2°", "E4 — Monitoraggio predittivo (Early Warning)",
         "Mercato lo chiede esplicitamente. Riutilizza M5 come base. Upsell naturale."],
        ["🟡 3°", "E6 — Cluster omogenei per early warning",
         "Si integra con E4. Differenziatore forte vs competitor."],
        ["🟡 4°", "E1 — Simulatore prodotto finanziario",
         "Sposta Metis da 'analisi' a 'raccomandazione commerciale'. Cambio di valore percepito elevato."],
        ["🟢 5°", "E8 — Verifica aderenza policy creditizie",
         "Valore compliance elevato. Si integra con Audit Trail esistente."],
        ["🟢 6°", "E2 — Indicatori di pagamento per cluster",
         "Utile ma richiede dati storici significativi per essere affidabile."],
        ["⚪ 7°", "E5 — Rischio AR",
         "Da chiarire il perimetro esatto prima di sviluppare."],
        ["⚪ 8°", "E3 — Teoria dei giochi",
         "Innovativa ma complessità molto alta. Da considerare solo dopo consolidamento core."],
    ]
)

# ══════════════════════════════════════════════════════════════════
# 7. CONCLUSIONE
# ══════════════════════════════════════════════════════════════════
heading1("7. Conclusione")

colorbox(
    "Metis ACE implementa il 100% dei requisiti core descritti nel documento Finomnia AI in PEF V05.\n\n"
    "Il documento Finomnia è stato la specifica funzionale che ha guidato lo sviluppo di Metis M1–M8.\n\n"
    "Esistono 8 evoluzioni future elencate nel documento che rappresentano la roadmap naturale del prodotto.\n\n"
    "Metis ha già aggiunto 3 funzionalità extra (Visual Policy Builder, Audit Trail, Fraud Detection) "
    "non presenti nella specifica originale, dimostrando che il prodotto ha già superato la visione iniziale.",
    fill_hex='D9EAD3', text_color=(0, 80, 0)
)

out = r"c:\Users\GaetanoPecorella\OneDrive - Altermaind\Desktop\METIS\Metis_GapAnalysis_Finomnia.docx"
doc.save(out)

import shutil
shutil.copy(out, r"C:\Users\GaetanoPecorella\Downloads\Metis_GapAnalysis_Finomnia.docx")
print(f"Salvato: {out}")
print("Copiato in Downloads.")

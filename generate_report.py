from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

doc = Document()

# ── Page margins ──────────────────────────────────────────────────────────────
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers ───────────────────────────────────────────────────────────────────
def set_font(run, bold=False, size=11, color=None, italic=False):
    run.bold   = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = "Calibri"
    if color:
        run.font.color.rgb = RGBColor(*color)

def heading1(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    set_font(run, bold=True, size=16, color=(0, 70, 127))
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    # bottom border
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'),  '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '0046FF')
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p

def heading2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, bold=True, size=13, color=(31, 73, 125))
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after  = Pt(4)
    return p

def body(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    set_font(run, bold=bold, italic=italic, size=11, color=color)
    p.paragraph_format.space_after = Pt(4)
    return p

def bullet(text, level=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    set_font(run, size=11)
    p.paragraph_format.space_after = Pt(2)
    return p

def add_table(headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = h
        cell.paragraphs[0].runs[0].bold = True
        cell.paragraphs[0].runs[0].font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].runs[0].font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        # blue background
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), '1F497D')
        tcPr.append(shd)
    # data rows
    for ri, row_data in enumerate(rows):
        row = table.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = val
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            if ri % 2 == 0:
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear')
                shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), 'DCE6F1')
                tcPr.append(shd)
    doc.add_paragraph()
    return table

# ══════════════════════════════════════════════════════════════════════════════
#  COVER
# ══════════════════════════════════════════════════════════════════════════════
cover = doc.add_paragraph()
cover.alignment = WD_ALIGN_PARAGRAPH.CENTER
cover.paragraph_format.space_before = Pt(60)
r = cover.add_run("METIS — McKinsey Strategic Analysis")
set_font(r, bold=True, size=24, color=(0, 70, 127))

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = sub.add_run("NanoBanana Pro: Posizionamento, Funzionalità & Competitor")
set_font(r2, italic=True, size=14, color=(89, 89, 89))

doc.add_paragraph()
date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = date_p.add_run("Marzo 2026  |  Confidenziale")
set_font(r3, size=11, color=(128, 128, 128))

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════════════════
#  1. EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════════════════════
heading1("1. Executive Summary")
body(
    "Metis è un motore AI Glass-Box per il credit underwriting di PMI italiane. "
    "Sostituisce un processo manuale da 6–9 ore con un'analisi completa in meno di 30 secondi, "
    "mantenendo piena compliance con l'EU AI Act e tracciabilità XAI (eXplainable AI) su ogni output.",
)
body(
    "NanoBanana Pro rappresenta il profilo-tipo del cliente Metis: una PMI italiana con bilancio XBRL, "
    "storico in Centrale Rischi e necessità di accedere rapidamente al credito bancario.",
)

# ══════════════════════════════════════════════════════════════════════════════
#  2. POSIZIONAMENTO
# ══════════════════════════════════════════════════════════════════════════════
heading1("2. Posizionamento nel Mercato — Mappa 2×2 McKinsey")
body(
    "Metis si colloca nel quadrante \"Automated + Explainable\" — il segmento meno affollato, "
    "più difendibile e allineato alle direttive EU AI Act (obbligatorie dal 2026).",
    bold=True
)
body("I quattro quadranti del mercato AI credit underwriting:")
bullet("Alta Explainability + Processo Automatizzato → METIS  ✅ (quadrante target)")
bullet("Alta Explainability + Processo Manuale → FICO Platform, Experian PowerCurve")
bullet("Bassa Explainability + Processo Automatizzato → Zest AI, Scienaptic AI, ACTICO")
bullet("Bassa Explainability + Processo Manuale → Processi bancari tradizionali")

# ══════════════════════════════════════════════════════════════════════════════
#  3. FUNZIONALITÀ
# ══════════════════════════════════════════════════════════════════════════════
heading1("3. Funzionalità Chiave — 8 Moduli ACE")

add_table(
    ["#", "Modulo", "Cosa fa", "Differenziatore"],
    [
        ["M1", "Narrative PEF",    "Genera testo analitico con link alle fonti",       "XAI source tracing — ogni frase è auditabile"],
        ["M2", "Web Sentiment",    "Crawling live + NLP keyword scoring",               "Nowcasting reputazionale real-time"],
        ["M3", "Financial KPIs",   "DSCR, EBITDA, Altman Z-Score",                     "Math deterministico (no LLM) = 100% affidabile"],
        ["M4", "ATECO Benchmark",  "Confronto con medie ISTAT di settore",             "Contesto settoriale per le banche"],
        ["M5", "CR Pattern 12M",   "Trend utilizzo Centrale Rischi",                   "Sparkline 12 mesi — anomalie visibili"],
        ["M6", "Cross-Check",      "Mismatch Bilancio vs CR",                          "Semaforo automatico per frodi documentali"],
        ["M7", "DSCR Forecast",    "Proiezione prospettica del debito",                "Forward-looking — raro tra i competitor"],
        ["M8", "SWOT Matrix",      "Generazione automatica 2×2",                       "Sintesi executive per il credit committee"],
    ]
)

body(
    "+ Visual Policy Builder: configurazione drag & drop del flusso decisionale (no-code) — unico nel segmento bancario italiano.",
    bold=True, italic=True
)

# ══════════════════════════════════════════════════════════════════════════════
#  4. COMPETITOR
# ══════════════════════════════════════════════════════════════════════════════
heading1("4. Landscape Competitivo")

heading2("4a. Competitor Diretti (Credit Underwriting AI)")
add_table(
    ["Player", "Sede", "Explainability", "EU AI Act", "Gap vs Metis"],
    [
        ["ACTICO",             "🇩🇪 Germania", "Parziale",  "Non dichiarato", "No Visual Policy Builder"],
        ["Zest AI",            "🇺🇸 USA",      "Moderata",  "Non conforme",   "No integrazione PEF italiano"],
        ["Scienaptic AI",      "🇺🇸 USA",      "Bassa",     "Non conforme",   "Black-box, no XAI"],
        ["Experian PowerCurve","🇬🇧 UK",       "Parziale",  "Parziale",       "Prodotto generico, non verticale"],
        ["FICO Platform",      "🇺🇸 USA",      "Alta",      "Parziale",       "Costoso, non adatto a banche medie IT"],
        ["RocketFin",          "🇮🇹 Italia",   "Moderata",  "Parziale",       "Nessuna integrazione CR/XBRL nativa"],
        ["Axe Credit Portal",  "🇧🇪 Belgio",   "Bassa",     "Parziale",       "No agentic AI swarm"],
    ]
)

heading2("4b. Competitor Indiretti")
add_table(
    ["Player", "Livello di Minaccia", "Note"],
    [
        ["Cerved (Analytics)",      "Media", "Fornisce i dati, potrebbe verticalizzare"],
        ["TeamSystem / Zucchetti",  "Alta",  "Già nel workflow bancario IT, potrebbero aggiungere AI"],
        ["Qonto / Finom",           "Bassa", "Focus lending diretto, non B2B verso banche"],
        ["Banche in-house (Intesa, UCG)", "Media", "Alcune banche stanno buildando internamente"],
    ]
)

# ══════════════════════════════════════════════════════════════════════════════
#  5. USP
# ══════════════════════════════════════════════════════════════════════════════
heading1("5. Unique Selling Proposition (USP)")
body("Metis è l'unico sistema che è contemporaneamente:", bold=True)
bullet("✅  Glass-Box (XAI su ogni output)")
bullet("✅  Verticale sul mercato bancario italiano (PEF, CR, XBRL)")
bullet("✅  EU AI Act compliant (human-in-the-loop obbligatorio)")
bullet("✅  No-code configurabile (Visual Policy Builder)")
bullet("✅  Agentic multi-agent (<30 secondi per pratica completa)")

# ══════════════════════════════════════════════════════════════════════════════
#  6. TAM / SAM / SOM
# ══════════════════════════════════════════════════════════════════════════════
heading1("6. Market Sizing — TAM / SAM / SOM")
add_table(
    ["Livello", "Definizione", "Stima"],
    [
        ["TAM", "Mercato globale AI credit decisioning", "~$12B (2025) → $38B (2030)"],
        ["SAM", "Banche & istituti finanziari in Europa", "~$2.4B"],
        ["SOM", "Banche medie-piccole italiane (target)", "~€120M"],
    ]
)
body("🎯 Target immediato: ~500 banche e istituti di credito in Italia con portafoglio PMI >€100M", bold=True)

# ══════════════════════════════════════════════════════════════════════════════
#  7. RACCOMANDAZIONI
# ══════════════════════════════════════════════════════════════════════════════
heading1("7. Raccomandazioni Strategiche")

heading2("A. Go-to-Market (Priorità Alta)")
bullet("Canale primario: Partnership con i principali CBS italiani — Temenos, CEDACRI, CSE")
bullet("Entry point: Pilota gratuito con 2–3 banche regionali (BCC, Casse Rurali) → poi espansione")

heading2("B. Posizionamento (Priorità Alta)")
bullet("Comunicare \"EU AI Act First\" come vantaggio competitivo — dal 2026 sarà un obbligo")
bullet("Naming: consolidare \"METIS ACE\" (Analisi Complementare Evoluta) come brand B2B")

heading2("C. Prodotto (Priorità Media)")
bullet("Accelerare M7 (DSCR Forecast) — raro tra competitor, alto valore percepito dalle banche")
bullet("Investire in connettori nativi: XBRL, API Bankitalia, API Cerved")
bullet("Visual Policy Builder → renderlo il principale differenziatore demo nel sales process")

heading2("D. Difesa Competitiva (Priorità Media)")
bullet("Depositare IP sul Visual Policy Builder + metodologia XAI")
bullet("Costruire moat sui dati: dataset storici PEF + benchmark ATECO proprietari")

# ══════════════════════════════════════════════════════════════════════════════
#  8. VERDICT
# ══════════════════════════════════════════════════════════════════════════════
heading1("8. Verdict")
body(
    "Metis si colloca nel segmento più strategico e meno affollato del mercato europeo del credit AI: "
    "quello \"regulated + explainable\". Con l'entrata in vigore dell'EU AI Act, i competitor black-box "
    "andranno incontro a significative frizioni normative. Metis è strutturalmente posizionato per "
    "catturare questa transizione.",
    italic=True
)
doc.add_paragraph()
add_table(
    ["Dimensione", "Rating"],
    [
        ["Attrattività del segmento",        "★★★★☆  (4/5)"],
        ["Vantaggio competitivo sostenibile", "★★★★★  (5/5 — se difeso con IP e partnership CBS)"],
    ]
)

# ── Save ─────────────────────────────────────────────────────────────────────
out = r"c:\Users\GaetanoPecorella\OneDrive - Altermaind\Desktop\METIS\Metis_McKinsey_Report.docx"
doc.save(out)
print(f"Salvato: {out}")

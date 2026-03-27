from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ── Page margins ──
for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# ── Helpers ──
def heading1(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(16); run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(0, 70, 127)
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after  = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bot = OxmlElement('w:bottom')
    bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), '6')
    bot.set(qn('w:space'), '1');    bot.set(qn('w:color'), '00467F')
    pBdr.append(bot); pPr.append(pBdr)

def heading2(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(13); run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(31, 73, 125)
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)

def heading3(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True; run.font.size = Pt(11); run.font.name = "Calibri"
    run.font.color.rgb = RGBColor(68, 114, 196)
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)

def body(text, bold=False, italic=False, color=None):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold; run.italic = italic
    run.font.size = Pt(11); run.font.name = "Calibri"
    if color: run.font.color.rgb = RGBColor(*color)
    p.paragraph_format.space_after = Pt(4)

def bullet(text, indent=0):
    p = doc.add_paragraph(style='List Bullet')
    run = p.add_run(text)
    run.font.size = Pt(11); run.font.name = "Calibri"
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(indent * 0.5)

def callout(text, color=(0xDC, 0xE6, 0xF1)):
    """Highlighted info box via a 1-cell table"""
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    cell.text = text
    run = cell.paragraphs[0].runs[0]
    run.font.size = Pt(11); run.font.name = "Calibri"; run.italic = True
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), '{:02X}{:02X}{:02X}'.format(*color))
    tcPr.append(shd)
    doc.add_paragraph()

def add_table(headers, rows, header_color='1F497D', alt_color='DCE6F1'):
    tbl = doc.add_table(rows=1 + len(rows), cols=len(headers))
    tbl.style = 'Table Grid'
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True; run.font.size = Pt(10); run.font.name = "Calibri"
        run.font.color.rgb = RGBColor(255, 255, 255)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        tc = cell._tc; tcPr = tc.get_or_add_tcPr()
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), header_color); tcPr.append(shd)
    for ri, row_data in enumerate(rows):
        row = tbl.rows[ri + 1]
        for ci, val in enumerate(row_data):
            cell = row.cells[ci]
            cell.text = val
            run = cell.paragraphs[0].runs[0]
            run.font.size = Pt(10); run.font.name = "Calibri"
            if ri % 2 == 0:
                tc = cell._tc; tcPr = tc.get_or_add_tcPr()
                shd = OxmlElement('w:shd')
                shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
                shd.set(qn('w:fill'), alt_color); tcPr.append(shd)
    doc.add_paragraph()

# ══════════════════════════════════════════════════════════════════
# COVER
# ══════════════════════════════════════════════════════════════════
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(60)
r = p.add_run("METIS ACE — Analisi Complementare Evoluta")
r.bold = True; r.font.size = Pt(24); r.font.name = "Calibri"
r.font.color.rgb = RGBColor(0, 70, 127)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p2.add_run("Destinatari, Processi e Funzionalità del Prodotto")
r2.italic = True; r2.font.size = Pt(14); r2.font.name = "Calibri"
r2.font.color.rgb = RGBColor(89, 89, 89)

doc.add_paragraph()
p3 = doc.add_paragraph()
p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
r3 = p3.add_run("Documento interno  |  Marzo 2026  |  Confidenziale")
r3.font.size = Pt(11); r3.font.name = "Calibri"
r3.font.color.rgb = RGBColor(128, 128, 128)

doc.add_page_break()

# ══════════════════════════════════════════════════════════════════
# 1. CHI È METIS
# ══════════════════════════════════════════════════════════════════
heading1("1. Cos'è Metis ACE")
body(
    "Metis è un motore AI Glass-Box per l'underwriting del credito alle PMI italiane. "
    "Automatizza la Pratica Elettronica di Fido (PEF) — il documento che ogni banca italiana "
    "produce per istruire una richiesta di affidamento — riducendo il tempo di analisi "
    "da 6–9 ore a meno di 30 secondi."
)
body(
    "A differenza dei sistemi black-box, Metis è Glass-Box: ogni output AI è tracciato "
    "alla fonte, ogni calcolo finanziario è deterministico (non LLM), e la decisione finale "
    "rimane sempre in mano all'analista umano (EU AI Act compliant)."
)

callout(
    "🔑  Principio fondante: 'AI come assistente, umano come decisore.' "
    "Metis non approva né rifiuta pratiche — le istruisce in modo completo, spiegabile e auditabile.",
    color=(0xD9, 0xEA, 0xD3)
)

# ══════════════════════════════════════════════════════════════════
# 2. DESTINATARI
# ══════════════════════════════════════════════════════════════════
heading1("2. Destinatari del Prodotto")

heading2("2.1 Utenti Primari — Gli Analisti del Credito")
body(
    "L'utente quotidiano di Metis è il Credit Analyst (o Gestore Imprese) della banca. "
    "Opera all'interno del CBS (Core Banking System) e ha la responsabilità di istruire "
    "le pratiche di fido per clienti PMI e Corporate."
)
add_table(
    ["Profilo", "Ruolo", "Obiettivo con Metis", "Pain point attuale"],
    [
        ["Analista Credito Senior", "Istruttoria pratiche complesse >500K€", "Ridurre il tempo di analisi, aumentare qualità", "6–9 ore per pratica, errori manuali, fonti eterogenee"],
        ["Analista Credito Junior", "Pratiche standard PMI <500K€", "Guida strutturata nell'analisi", "Mancanza di esperienza su KPI e benchmark settoriali"],
        ["Gestore Imprese (RM)", "Relazione con il cliente", "Avere una sintesi pronta prima del colloquio", "Non ha tempo per analisi quantitativa approfondita"],
        ["Credit Officer / Risk Manager", "Supervisione e firma del deliberato", "Audit trail completo, conformità EU AI Act", "Difficoltà nel controllare la qualità dei dossier"],
    ]
)

heading2("2.2 Utenti Secondari — Decision Makers")
bullet("Direttore Crediti: supervisiona portafoglio e delibere — usa Metis per monitoraggio KPI aggregati")
bullet("Compliance Officer: verifica rispetto EU AI Act — usa l'Audit Trail immutabile di Metis")
bullet("CRO (Chief Risk Officer): oversight del rischio di portafoglio — dashboard aggregata")

heading2("2.3 Cliente Finale (Indiretto)")
body(
    "La PMI italiana che richiede un finanziamento. Beneficia indirettamente di Metis attraverso "
    "risposte più rapide dalla banca (da settimane a ore) e valutazioni più accurate e meno "
    "discriminatorie (grazie al modello FairBoost XGBoost)."
)
add_table(
    ["Profilo PMI", "Settore tipo", "Dati forniti a Metis"],
    [
        ["Micro PMI (1–9 dipendenti)", "Artigianato, commercio al dettaglio", "Bilancio XBRL, 1 anno CR"],
        ["PMI Media (10–49 dipendenti)", "Manifatturiero, servizi B2B", "Bilancio XBRL, 3 anni CR, storico PEF"],
        ["PMI Grande (50–249 dipendenti)", "Industria, export, costruzioni", "Bilancio XBRL consolidato, CR completo, Cerved"],
    ]
)

# ══════════════════════════════════════════════════════════════════
# 3. IL PROCESSO — COME FUNZIONA
# ══════════════════════════════════════════════════════════════════
heading1("3. Il Processo — Come Funziona Metis")

heading2("3.1 Flusso End-to-End")
body("Il processo Metis si articola in 7 fasi sequenziali, dalla richiesta all'approvazione:")

add_table(
    ["Fase", "Attore", "Azione", "Output", "Tempo"],
    [
        ["1. Trigger",          "Analista",  "Clicca 'Avvia ACE' nel CBS o nel pannello Metis",          "Job creato in coda",                          "< 1 sec"],
        ["2. Data Collection",  "Sistema",   "Raccoglie dati da Bilancio XBRL, CR, PEF storici, Web",    "Dati staged + quality score per fonte",       "5–15 sec"],
        ["3. Validazione",      "Sistema",   "Controlla completezza e qualità dei dati (Pydantic)",       "✅ Processa / ⚠️ Richiede revisione analista", "< 2 sec"],
        ["4. Analisi (8 mod.)", "Sistema AI","Esegue M1–M8 in parallelo (DAG Prefect)",                  "8 output JSON + spiegazioni XAI (SHAP)",       "10–20 sec"],
        ["5. Report",           "Sistema",   "Assembla output in PEF narrative completa",                 "PDF + HTML scaricabili",                      "< 5 sec"],
        ["6. Human Review",     "Analista",  "Legge, verifica, eventualmente corregge ogni sezione",      "Conferma / Override / Escalation",            "20–30 min"],
        ["7. Delibera",         "Analista",  "Approva / Richiede integrazioni / Rifiuta",                 "Decisione finale + audit record immutabile",  "5–10 min"],
    ]
)

callout(
    "⏱  Risparmio di tempo stimato: da 6–9 ore a ~30 minuti di revisione umana. "
    "La differenza è che Metis produce il dossier grezzo completo; l'analista "
    "si concentra solo sulla verifica e sulla decisione.",
    color=(0xFF, 0xF2, 0xCC)
)

heading2("3.2 Fonti Dati Aggregate")
add_table(
    ["Fonte", "Tipo", "Moduli che la usano", "Note"],
    [
        ["Bilancio XBRL",        "Strutturato",     "M3, M6, M7, M8",   "Formato standard europeo — parser deterministico"],
        ["Centrale Rischi (CR)", "Strutturato",     "M5, M6, M7, M8",   "12 mesi di storico mensile, TimescaleDB"],
        ["Commenti PEF storici", "Testo libero",    "M1",               "Commenti analisti precedenti — input LLM"],
        ["Web / DuckDuckGo",     "Non strutturato", "M2",               "News, Registro Imprese, Il Sole 24 Ore, ANSA"],
        ["ISTAT / ATECO",        "Statistico",      "M4, M7, M8",       "Mediane settoriali per benchmark"],
        ["Cerved APIs",          "Strutturato",     "Data Ingestion",   "Dati camerali, protesti, rating Cerved"],
    ]
)

heading2("3.3 Visual Policy Builder — Il Controllo del Flusso Decisionale")
body(
    "Oltre al motore AI, Metis include un editor grafico drag & drop (React Flow) "
    "che permette al Credit Officer o al Risk Manager di configurare le regole "
    "decisionali senza scrivere codice."
)
bullet("I 'blocchi' disponibili: Data Ingestion, Fraud Detection, AI Scoring, Auto-Approve, Reject/Escalate, Custom Rule")
bullet("Ogni blocco ha parametri configurabili: soglie di score, pesi delle fonti, limiti di falsi positivi")
bullet("Il flusso viene salvato come policy JSON e versioned — ogni modifica è tracciata nell'Audit Trail")
bullet("Esempio NanoBanana Pro: Score > 65 → Auto-Approve con signoff | Score < 35 → Reject automatico")

# ══════════════════════════════════════════════════════════════════
# 4. FUNZIONALITÀ — GLI 8 MODULI ACE
# ══════════════════════════════════════════════════════════════════
heading1("4. Funzionalità — Gli 8 Moduli ACE")

body(
    "Il cuore di Metis sono gli 8 moduli indipendenti (M1–M8), ciascuno con il proprio "
    "scope, tecnologia e output. Possono essere attivati singolarmente o tutti insieme "
    "in sequenza orchestrata."
)

# M1
heading2("M1 — Sintesi Commenti Storici PEF")
heading3("Cosa fa")
body(
    "Legge i commenti testuali scritti dagli analisti nelle pratiche PEF precedenti "
    "dello stesso cliente (fino a 3 annualità) e produce una sintesi narrativa strutturata, "
    "evidenziando le variazioni significative rispetto alle annualità precedenti (Delta Highlights)."
)
heading3("Come funziona")
bullet("Input: testo libero dei commenti storici (bilancio + CR + andamentale)")
bullet("Tecnologia: LLM (Claude Sonnet / GPT-4o) con prompt strutturato + validazione anti-allucinazione")
bullet("Anti-hallucination: ogni numero nell'output viene verificato contro i numeri nel testo sorgente")
bullet("XAI: source attribution — ogni paragrafo è linkato alla frase sorgente da cui deriva")
heading3("Output")
bullet("sintesi_societaria: profilo dell'azienda in 3–5 righe")
bullet("sintesi_bilancio: andamento economico-patrimoniale sintetizzato")
bullet("sintesi_cr: comportamento bancario storico sintetizzato")
bullet("delta_highlights: variazioni significative anno su anno con flag di rilevanza (HIGH/MEDIUM/LOW)")
bullet("confidence_scores: confidenza dell'LLM per ogni paragrafo prodotto")
doc.add_paragraph()

# M2
heading2("M2 — Sintesi Commenti da Fonti Web (Nowcasting Reputazionale)")
heading3("Cosa fa")
body(
    "Cerca notizie e informazioni pubbliche sull'azienda su fonti whitelistate "
    "(Il Sole 24 Ore, ANSA, Reuters Italia, Registro Imprese) e produce un'analisi "
    "del sentiment reputazionale aggiornata al giorno dell'analisi."
)
heading3("Come funziona")
bullet("Web scraping con rate limiting (2 req/sec) su fonti approvate — no hallucination risk")
bullet("Sentiment analysis: modello FinBERT specializzato in linguaggio finanziario italiano")
bullet("Aggregazione sentiment: media pesata (fonti più autorevoli = peso maggiore)")
bullet("Sintesi narrativa finale via LLM sull'aggregato degli articoli trovati")
heading3("Output")
bullet("descrizione_attivita: descrizione aggiornata dell'attività dell'azienda da fonti web")
bullet("sentiment_complessivo: POSITIVO / NEUTRO / NEGATIVO con score 0–1")
bullet("fonti: lista articoli con URL, data, estratto e sentiment individuale")
bullet("soggetti_correlati: persone o aziende correlate emerse dalle notizie")
doc.add_paragraph()

# M3
heading2("M3 — Analisi KPI Bilancio (Financial Health Score)")
heading3("Cosa fa")
body(
    "Calcola in modo deterministico (pura matematica Python, zero LLM) i 9 KPI "
    "finanziari fondamentali confrontando due esercizi (T0 vs T1), genera alert "
    "sulle soglie e produce uno score di rischio XGBoost con spiegazione SHAP."
)
heading3("KPI Calcolati")
add_table(
    ["KPI", "Formula", "Soglia Alert"],
    [
        ["EBITDA Margin",       "EBITDA / Ricavi",                           "< 5%"],
        ["ROE",                 "Utile Netto / Patrimonio Netto",             "< 2%"],
        ["ROA",                 "Utile Netto / Attivo Totale",               "< 1%"],
        ["Leverage (D/E)",      "Debiti Fin. / Patrimonio Netto",            "> 4x"],
        ["Current Ratio",       "Attivo Corrente / Passivo Corrente",        "< 1.0"],
        ["Interest Coverage",   "EBITDA / Oneri Finanziari",                 "< 1.5"],
        ["NFP/EBITDA",          "(Debiti Fin. - Liquidità) / EBITDA",       "> 5x"],
        ["DSO",                 "Crediti Comm. / Ricavi × 365",              "> 120 gg"],
        ["DPO",                 "Debiti Comm. / Costi Op. × 365",           "> 180 gg"],
    ]
)
heading3("Output aggiuntivi")
bullet("Trend per ogni KPI: UP / STABLE / DOWN con variazione percentuale")
bullet("Risk Score 0–1 (XGBoost) con feature importance SHAP")
bullet("Altman Z-Score calcolato deterministicamente")
doc.add_paragraph()

# M4
heading2("M4 — Benchmark ATECO / ISTAT (Posizionamento Settoriale)")
heading3("Cosa fa")
body(
    "Confronta i KPI dell'azienda analizzata con le mediane settoriali ISTAT "
    "per codice ATECO, determinando se l'azienda è OUTPERFORMER, IN LINEA o UNDERPERFORMER."
)
heading3("Come funziona")
bullet("Database ISTAT: mediane, percentili 25° e 75° per ~200 codici ATECO")
bullet("Matching automatico: codice ATECO dell'azienda → gruppo settoriale corrispondente")
bullet("Calcolo percentile: posizionamento dell'azienda nella distribuzione settoriale")
heading3("Output")
bullet("Posizione per ogni KPI: OUTPERFORMER (>75° percentile) / IN LINEA / UNDERPERFORMER (<25°)")
bullet("Narrativa: 'L'azienda mostra un EBITDA Margin del 12%, superiore alla mediana settoriale (8%)'")
bullet("Input per M7 (Forecast) e M8 (SWOT)")
doc.add_paragraph()

# M5
heading2("M5 — Analisi CR 12 Mesi (Pattern Centrale Rischi)")
heading3("Cosa fa")
body(
    "Analizza 12 mesi di dati mensili della Centrale Rischi (accordato, utilizzato, "
    "sconfinamenti, sofferenze, scaduto) usando analisi statistica e anomaly detection "
    "per identificare pattern di rischio e comportamenti anomali."
)
heading3("Come funziona")
bullet("Time-series analysis: calcolo medie, delta, utilizzo medio su 12 mesi")
bullet("Anomaly detection: Z-score statistico — valori > 2σ vengono flaggati come anomali")
bullet("Classificazione anomalie: TRANSITORIA (rientrata entro 60 giorni) vs STRUTTURALE (persistente)")
bullet("Utilizzo ratio: accordato vs utilizzato — alert sopra 85% (rischio saturazione linee)")
heading3("Output")
bullet("metriche: tabella 12 mesi con valori, media, delta e classificazione (OK/WARNING/CRITICAL)")
bullet("anomalie: lista eventi anomali con Z-score, mese, valore atteso vs rilevato")
bullet("risk_score_cr: score 0–1 di rischio CR con spiegazione XAI")
bullet("narrativa: testo descrittivo del comportamento CR in linguaggio bancario")
doc.add_paragraph()

# M6
heading2("M6 — Cross-Check CR vs Bilancio (Rilevamento Incoerenze)")
heading3("Cosa fa")
body(
    "Confronta i debiti bancari dichiarati nel bilancio con i dati reali della "
    "Centrale Rischi, rilevando eventuali disallineamenti che potrebbero indicare "
    "omissioni, errori contabili o manipolazioni."
)
heading3("Come funziona")
bullet("Calcolo delta percentuale: |Bilancio - CR| / max(Bilancio, CR) × 100")
bullet("Soglie: 0–10% → OK | 10–25% → WARNING | >25% → CRITICAL")
bullet("Analisi cause automatica: sfasamento temporale, debiti non censiti in CR, utilizzi post-chiusura")
heading3("Output")
bullet("delta_pct_totale: scarto percentuale totale tra bilancio e CR")
bullet("severity: OK / WARNING / CRITICAL con semaforo visivo")
bullet("possibili_cause: lista ragioni plausibili dello scostamento")
bullet("raccomandazione: azione suggerita all'analista (es. 'Richiedere documentazione integrativa')")
doc.add_paragraph()

# M7
heading2("M7 — Forecast DSCR Prospettico (Sostenibilità del Debito)")
heading3("Cosa fa")
body(
    "Calcola il Debt Service Coverage Ratio (DSCR) prospettico su 3 scenari "
    "(Ottimistico / Base / Stress) per valutare la capacità futura dell'azienda "
    "di servire il debito richiesto."
)
heading3("Formula DSCR")
callout(
    "DSCR = (EBITDA proiettato − Tasse − Capex mantenimento) ÷ (Quota capitale annua + Interessi annui)\n\n"
    "• DSCR > 1.3  →  SOSTENIBILE\n"
    "• 1.0 < DSCR < 1.3  →  BORDERLINE\n"
    "• DSCR < 1.0  →  NON SOSTENIBILE"
)
heading3("I 3 Scenari")
add_table(
    ["Scenario", "Crescita Ricavi", "Crescita Costi", "Delta Tassi", "Quando si applica"],
    [
        ["OTTIMISTICO", "+10%", "+3%", "0%",    "CR OK + KPI in crescita"],
        ["BASE",        "0%",   "+2%", "+0.5%", "Situazione standard"],
        ["STRESS",      "-15%", "+5%", "+1.5%", "CR CRITICAL o KPI in calo"],
    ]
)
body("Lo scenario viene selezionato automaticamente dai dati di M3 e M5, ma l'analista può forzarlo.")
doc.add_paragraph()

# M8
heading2("M8 — SWOT Matrix (Sintesi Strategica)")
heading3("Cosa fa")
body(
    "Aggrega automaticamente gli output di M2, M3, M4 e M5 in una matrice SWOT "
    "strutturata, distinguendo fattori interni (Forze/Debolezze) da fattori esterni "
    "(Opportunità/Minacce), con evidenza numerica per ogni punto."
)
heading3("Logica di mapping")
bullet("FORZE: KPI sopra mediana ISTAT + CR regolare + sentiment web positivo")
bullet("DEBOLEZZE: KPI in alert + anomalie CR strutturali + trend negativo")
bullet("OPPORTUNITÀ: trend settoriale positivo + indicatori macro favorevoli")
bullet("MINACCE: sentiment web negativo + forecast DSCR Stress < 1.0 + esposizione a tassi variabili")
heading3("Output")
bullet("Matrice 2×2 con 3–5 item per quadrante, ciascuno con evidenza numerica e modulo sorgente")
bullet("Ogni item è linkato al modulo che lo ha generato (source_module: M3, M4, M5...)")
bullet("Narrativa LLM che sintetizza i 4 quadranti in un paragrafo executive")

# ══════════════════════════════════════════════════════════════════
# 5. FUNZIONALITÀ TRASVERSALI
# ══════════════════════════════════════════════════════════════════
heading1("5. Funzionalità Trasversali")

heading2("5.1 XAI — Explainable AI")
body(
    "Ogni modulo che produce uno score o una predizione genera automaticamente "
    "una spiegazione tramite SHAP (SHapley Additive exPlanations)."
)
bullet("Feature importance: quali KPI hanno pesato di più nello score (es. 'il Leverage ha contribuito +0.23 al rischio')")
bullet("Source attribution: ogni paragrafo LLM è linkato alla frase sorgente che lo ha generato")
bullet("Confidence scores: grado di certezza dell'LLM per ogni sezione prodotta")

heading2("5.2 Audit Trail Immutabile")
body(
    "Ogni azione nel sistema — esecuzione di un modulo, conferma umana, override, escalation — "
    "genera un record immutabile nel database con hash SHA-256 degli input e output."
)
bullet("Chi ha fatto cosa e quando: actor, timestamp, job_id, module_code")
bullet("Override obbligatorio con motivazione: se l'analista corregge un output AI, deve spiegare perché")
bullet("Exportable: il trail è consultabile dal Compliance Officer e scaricabile in PDF")

heading2("5.3 Data Quality Scoring")
body("Prima di elaborare, ogni fonte dati riceve uno score di qualità 0–100.")
add_table(
    ["Fonte", "Quality Score 85%", "Quality Score < 50% → Azione"],
    [
        ["Bilancio XBRL",  "Procede normalmente",    "Warning a analista — possibili dati mancanti"],
        ["CR Data",        "Procede normalmente",    "Job bloccato — richiesta integrazione obbligatoria"],
        ["Web Sources",    "Procede normalmente",    "Modulo M2 in modalità 'parziale' — dati limitati"],
    ]
)

heading2("5.4 Report Generator")
bullet("Output: PDF e HTML scaricabili con tutti gli 8 moduli assemblati in PEF narrativa completa")
bullet("Stile bancario standard — compatibile con il formato di delibera delle banche italiane")
bullet("Include disclaimer EU AI Act obbligatorio in ogni report")

heading2("5.5 Visual Policy Builder (No-Code)")
bullet("Editor drag & drop per configurare le regole decisionali del processo di credito")
bullet("Blocchi disponibili: Data Ingestion, Fraud Detection, AI Scoring, Auto-Approve, Reject, Custom Rule")
bullet("Ogni policy è versionata e storicizzata — roll-back possibile a qualsiasi versione precedente")
bullet("Permette alla banca di personalizzare soglie e pesi senza modificare il codice")

# ══════════════════════════════════════════════════════════════════
# 6. COMPLIANCE EU AI ACT
# ══════════════════════════════════════════════════════════════════
heading1("6. Compliance EU AI Act")
body(
    "Metis è progettato nativamente per soddisfare i requisiti dell'EU AI Act "
    "(Regolamento UE 2024/1689), che classifica i sistemi AI per il credit scoring "
    "come 'Alto Rischio' e impone obblighi specifici."
)
add_table(
    ["Requisito EU AI Act", "Come Metis lo soddisfa"],
    [
        ["Human oversight obbligatorio",         "Delibera Action Bar: Approva / Richiede Integrazioni / Rifiuta — l'analista decide sempre"],
        ["Trasparenza e spiegabilità",            "XAI SHAP su ogni score + source attribution LLM su ogni paragrafo"],
        ["Audit trail",                           "Ogni azione registrata in tabella audit_trail con hash SHA-256 immutabile"],
        ["Nessuna decisione automatica definitiva","Il sistema non approva né rifiuta — produce solo una proposta analitica"],
        ["Accuracy e robustezza",                 "Math deterministico per i calcoli finanziari — zero LLM per i KPI"],
        ["Data governance",                       "Dati sintetici in sviluppo, nessun dato reale senza consenso"],
    ]
)

# ══════════════════════════════════════════════════════════════════
# 7. STACK TECNOLOGICO
# ══════════════════════════════════════════════════════════════════
heading1("7. Stack Tecnologico")
add_table(
    ["Layer", "Tecnologia", "Scopo"],
    [
        ["Frontend",       "Next.js 16 + TailwindCSS v4 + React Flow", "UI analista + Visual Policy Builder"],
        ["Backend API",    "FastAPI (Python)",                          "Orchestratore agenti + endpoint REST"],
        ["AI / LLM",       "Claude Sonnet (primario) + GPT-4o (fallback)", "Moduli narrativi M1, M2, M7, M8"],
        ["ML / Scoring",   "XGBoost + LightGBM + scikit-learn",        "Risk scoring M3, M5"],
        ["XAI",            "SHAP (TreeExplainer + KernelExplainer)",   "Spiegabilità su ogni prediction"],
        ["Database",       "PostgreSQL 16 + TimescaleDB",              "Dati strutturati + time-series CR"],
        ["Cache / Queue",  "Redis + RabbitMQ",                         "Job queue e caching"],
        ["Orchestrazione", "Prefect",                                   "DAG per esecuzione moduli in parallelo"],
        ["Web Scraping",   "DuckDuckGo API + BeautifulSoup",           "Modulo M2 — nowcasting reputazionale"],
    ]
)

# ── Save ──
out = r"c:\Users\GaetanoPecorella\OneDrive - Altermaind\Desktop\METIS\Metis_Prodotto_Completo.docx"
doc.save(out)
print(f"Salvato: {out}")
